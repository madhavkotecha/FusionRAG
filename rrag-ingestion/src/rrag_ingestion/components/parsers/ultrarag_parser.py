from __future__ import annotations

import logging
import os
import uuid
from datetime import datetime
from typing import Any

from rrag_ingestion.components.parsers.base import BaseParser
from rrag_ingestion.core.base import ComponentResult, register_component
from rrag_ingestion.models.document import Document

logger = logging.getLogger(__name__)


@register_component
class UltraRAGMultiParser(BaseParser):
    name = "ultrarag.multi_parser"
    version = "0.1.0"
    description = "Multi-format document parser (PDF, DOCX, MD, TXT, EPUB) extracted from UltraRAG"
    config_schema = {
        "formats": {"type": "array", "items": {"type": "string"}, "default": ["pdf", "docx", "md", "txt"]},
        "encoding": {"type": "string", "default": "utf-8"},
    }

    async def invoke(self, input_data: ComponentResult, config: dict) -> ComponentResult:
        file_paths = input_data.data.get("file_paths", [])
        existing_docs = input_data.data.get("documents", [])
        encoding = config.get("encoding", "utf-8")
        allowed_formats = set(config.get("formats", ["pdf", "docx", "md", "txt"]))

        documents = list(existing_docs)
        errors = []

        for fi, fp in enumerate(file_paths):
            self.report_progress(
                f"Parsing file {fi + 1}/{len(file_paths)}: {os.path.basename(fp)}",
                items_done=fi,
                items_total=len(file_paths),
            )
            ext = os.path.splitext(fp)[1].lower().lstrip(".")
            if ext not in allowed_formats:
                errors.append(f"Skipped unsupported format: {fp}")
                continue

            try:
                raw_text = await self._parse_file(fp, ext, encoding)
                doc = Document(
                    id=str(uuid.uuid4()),
                    filename=os.path.basename(fp),
                    content_type=self._mime_type(ext),
                    raw_text=raw_text,
                    file_path=fp,
                    metadata={"parser": self.name, "format": ext, "size_bytes": os.path.getsize(fp)},
                )
                documents.append(doc)
            except Exception as e:
                errors.append(f"Failed to parse {fp}: {e}")

        return ComponentResult(
            data={"documents": documents},
            metadata={"parsed_count": len(documents), "skipped_count": len(errors)},
            errors=errors,
        )

    async def _parse_file(self, file_path: str, ext: str, encoding: str) -> str:
        if ext == "pdf":
            return self._parse_pdf(file_path)
        elif ext == "docx":
            return self._parse_docx(file_path)
        else:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()

    def _parse_pdf(self, file_path: str) -> str:
        try:
            import fitz  # pymupdf
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            logger.warning("pymupdf not installed, reading PDF as binary placeholder")
            return f"[PDF content from {os.path.basename(file_path)} - install pymupdf for parsing]"

    def _parse_docx(self, file_path: str) -> str:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            logger.warning("python-docx not installed")
            return f"[DOCX content from {os.path.basename(file_path)} - install python-docx for parsing]"

    def _mime_type(self, ext: str) -> str:
        return {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "md": "text/markdown",
            "txt": "text/plain",
            "epub": "application/epub+zip",
        }.get(ext, "application/octet-stream")
